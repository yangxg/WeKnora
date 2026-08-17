"""Generate the real DOCX regression fixture for GitHub issue #2634."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

OUTPUT = Path(__file__).with_name("issue_2634_vertical_merge.docx")


def main() -> None:
    document = Document()
    document.add_heading("遗传病检测项目表", level=1)
    document.add_paragraph(
        "回归场景：检测方法纵向合并，但每条检测项目在转换后都必须保留该方法。"
    )

    rows = [
        ("Q0101", "遗传性乳腺癌", "BRCA1", "15个工作日"),
        ("Q0102", "遗传性卵巢癌", "BRCA2", "15个工作日"),
        ("Q0103", "林奇综合征", "MLH1", "15个工作日"),
        ("Q0104", "家族性腺瘤性息肉病", "APC", "15个工作日"),
    ]
    method = (
        "检测方法：提取外周血基因组 DNA，采用高通量测序，并对候选位点进行"
        "Sanger 测序验证；检测结果需结合临床表现综合判断。"
    )

    table = document.add_table(rows=len(rows) + 1, cols=5)
    table.style = "Table Grid"
    headers = ("项目编号", "检测项目", "相关基因", "周期", "检测方法")
    for column, value in enumerate(headers):
        table.cell(0, column).text = value

    for row_index, row in enumerate(rows, start=1):
        for column, value in enumerate(row):
            table.cell(row_index, column).text = value

    merged_method = table.cell(1, 4).merge(table.cell(len(rows), 4))
    merged_method.text = method
    merged_method.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_paragraph("非合并对照表：相邻单元格内容允许相同，但不应被当作合并。")
    control = document.add_table(rows=2, cols=3)
    control.style = "Table Grid"
    for column, value in enumerate(("列A", "列B", "列C")):
        control.cell(0, column).text = value
    for column, value in enumerate(("相同值", "相同值", "独立值")):
        control.cell(1, column).text = value

    document.save(OUTPUT)
    print(f"generated={OUTPUT}")
    print(f"size={OUTPUT.stat().st_size}")


if __name__ == "__main__":
    main()
